from typing import Dict, Any, Optional
import base64
import io
from datetime import datetime
import logging
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak, KeepTogether, HRFlowable
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.graphics.shapes import Drawing, Rect, Line
from reportlab.graphics import renderPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml import parse_xml
import PIL.Image

logger = logging.getLogger(__name__)

class ExportService:
    def __init__(self):
        self.styles = getSampleStyleSheet()
        self.custom_styles = self._create_custom_styles()
        self._register_fonts()

    def _register_fonts(self):
        """Register Roboto fonts if available, fallback to Helvetica"""
        try:
            # Try to register Roboto fonts (would need font files in production)
            # For now, we'll use Helvetica as fallback
            self.font_family = 'Helvetica'
            self.font_bold = 'Helvetica-Bold'
        except:
            self.font_family = 'Helvetica'
            self.font_bold = 'Helvetica-Bold'

    def _create_custom_styles(self):
        """Create custom styles for PDF generation with enhanced styling"""
        custom_styles = {}
        
        # Enhanced title style with gradient-like effect
        custom_styles['title'] = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=26,
            spaceAfter=20,
            spaceBefore=10,
            textColor=colors.HexColor('#1a365d'),  # Dark blue
            alignment=1,  # Center alignment
            fontName='Helvetica-Bold',
            borderWidth=2,
            borderPadding=12,
            borderColor=colors.HexColor('#3182ce'),
            backColor=colors.HexColor('#f7fafc')  # Light background
        )
        
        # Professional subtitle style
        custom_styles['subtitle'] = ParagraphStyle(
            'CustomSubtitle',
            parent=self.styles['Normal'],
            fontSize=14,
            spaceAfter=25,
            textColor=colors.HexColor('#4a5568'),  # Medium gray
            alignment=1,
            fontName='Helvetica-Oblique',
            borderWidth=1,
            borderPadding=8,
            borderColor=colors.HexColor('#e2e8f0'),
            backColor=colors.HexColor('#ffffff')
        )
        
        # Enhanced section header with professional gradient effect
        custom_styles['section'] = ParagraphStyle(
            'CustomSection',
            parent=self.styles['Heading2'],
            fontSize=16,
            spaceBefore=30,
            spaceAfter=15,
            textColor=colors.white,
            leftIndent=15,
            fontName='Helvetica-Bold',
            borderWidth=2,
            borderPadding=12,
            borderColor=colors.HexColor('#2d3748')
        )
        
        # Enhanced subsection style
        custom_styles['subsection'] = ParagraphStyle(
            'CustomSubsection',
            parent=self.styles['Heading3'],
            fontSize=14,
            spaceBefore=20,
            spaceAfter=12,
            textColor=colors.HexColor('#2d3748'),
            leftIndent=10,
            fontName='Helvetica-Bold',
            borderWidth=1,
            borderPadding=8,
            borderColor=colors.HexColor('#cbd5e0'),
            backColor=colors.HexColor('#f7fafc')
        )
        
        # Professional body text with subtle styling
        custom_styles['body'] = ParagraphStyle(
            'CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            leftIndent=25,
            fontName='Helvetica',
            textColor=colors.HexColor('#2d3748')
        )
        
        # Company info style
        custom_styles['company'] = ParagraphStyle(
            'CompanyInfo',
            fontSize=20,
            textColor=colors.HexColor('#1a365d'),
            fontName='Helvetica-Bold',
            alignment=1,
            spaceBefore=10,
            spaceAfter=5
        )
        
        # Contact info style
        custom_styles['contact'] = ParagraphStyle(
            'ContactInfo',
            fontSize=11,
            textColor=colors.HexColor('#4a5568'),
            fontName='Helvetica',
            alignment=1,
            spaceAfter=3
        )
        
        # Date style with enhanced formatting
        custom_styles['date'] = ParagraphStyle(
            'DateStyle',
            fontSize=11,
            textColor=colors.HexColor('#718096'),
            alignment=1,
            fontName='Helvetica-Oblique',
            borderWidth=1,
            borderPadding=6,
            borderColor=colors.HexColor('#e2e8f0'),
            backColor=colors.HexColor('#f7fafc')
        )
        
        return custom_styles

    def _create_section_header(self, title, bg_color):
        """Create a styled section header with enhanced background and borders"""
        header_style = ParagraphStyle(
            'SectionHeader',
            fontSize=16,
            spaceBefore=30,
            spaceAfter=15,
            textColor=colors.white,
            leftIndent=15,
            fontName='Helvetica-Bold',
            backColor=bg_color,
            borderPadding=12,
            borderWidth=2,
            borderColor=colors.HexColor('#2d3748')
        )
        return Paragraph(title, header_style)

    def _create_enhanced_divider_line(self, color=None, thickness=2):
        """Create an enhanced horizontal divider line with professional styling"""
        if color is None:
            color = colors.HexColor('#3182ce')
        return HRFlowable(width="100%", thickness=thickness, color=color, spaceBefore=10, spaceAfter=10)

    def _create_decorative_border(self):
        """Create a decorative border element"""
        return HRFlowable(width="100%", thickness=3, color=colors.HexColor('#e2e8f0'), spaceBefore=5, spaceAfter=5)

    def _create_logo_with_text_table(self, text_data, logo_data, logo_label):
        """Create a table that displays text with logo next to it"""
        table_data = []
        
        for text_item in text_data:
            label, value = text_item
            if label == logo_label and logo_data:
                # Create a table row with text and logo
                logo_img = self._get_logo_image(logo_data, height=0.8*inch, width=1.2*inch)
                if logo_img:
                    table_data.append([label, value, logo_img])
                else:
                    table_data.append([label, value, ""])
            else:
                table_data.append([label, value, ""])
        
        if table_data:
            # Use different column widths if logos are present
            has_logos = any(len(row) > 2 and row[2] != "" for row in table_data)
            if has_logos:
                col_widths = [2*inch, 2.5*inch, 1.5*inch]
                table = Table(table_data, colWidths=col_widths)
            else:
                # Standard two-column layout
                table_data = [[row[0], row[1]] for row in table_data]
                col_widths = [2.5*inch, 3.5*inch]
                table = Table(table_data, colWidths=col_widths)
            
            # Enhanced table styling with professional borders
            table_style = [
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 11),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 10),
                ('TOPPADDING', (0, 0), (-1, -1), 10),
                ('LEFTPADDING', (0, 0), (-1, -1), 15),
                ('RIGHTPADDING', (0, 0), (-1, -1), 15),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
                ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#3182ce')),  # Header underline
            ]
            
            # Special styling for logo cells
            if has_logos:
                table_style.extend([
                    ('ALIGN', (2, 0), (2, -1), 'CENTER'),
                    ('VALIGN', (2, 0), (2, -1), 'MIDDLE'),
                ])
            
            table.setStyle(TableStyle(table_style))
            return table
        return None

    def _should_include_field(self, value: Any) -> bool:
        """Check if field should be included (not empty)"""
        if value is None:
            return False
        if isinstance(value, str):
            return value.strip() != ''
        if isinstance(value, (list, dict)):
            return len(value) > 0
        return True

    def _get_logo_image(self, logo_data: str, height=1*inch, width=2*inch) -> Optional[Image]:
        """Convert base64 logo data to ReportLab Image with custom sizing"""
        try:
            if logo_data and logo_data.startswith('data:image'):
                # Extract base64 data
                base64_data = logo_data.split(',')[1]
                image_data = base64.b64decode(base64_data)
                image_buffer = io.BytesIO(image_data)
                
                # Create ReportLab Image with custom dimensions
                img = Image(image_buffer)
                img.drawHeight = height
                img.drawWidth = width
                img.hAlign = 'CENTER'
                return img
        except Exception as e:
            logger.error(f"Error processing logo: {str(e)}")
        return None

    def _get_logo_from_data(self, data: Dict[str, Any], logo_path: str) -> Optional[str]:
        """Extract logo data from nested dictionary path"""
        try:
            keys = logo_path.split('.')
            current = data
            for key in keys:
                if isinstance(current, dict) and key in current:
                    current = current[key]
                else:
                    return None
            
            # Handle both old format (string) and new format (dict with dataUrl)
            if isinstance(current, dict) and 'dataUrl' in current:
                return current['dataUrl']
            elif isinstance(current, str):
                return current
        except Exception as e:
            logger.error(f"Error extracting logo from {logo_path}: {str(e)}")
        return None

    def _create_styled_table(self, data, col_widths, bg_color=None, has_logos=False):
        """Create a styled table with enhanced formatting and logo support"""
        if not data:
            return None
            
        table = Table(data, colWidths=col_widths)
        
        # Enhanced table styling with professional appearance
        table_style = [
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 12),
            ('LEFTPADDING', (0, 0), (-1, -1), 15),
            ('RIGHTPADDING', (0, 0), (-1, -1), 15),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#cbd5e0')),
            ('ROWBACKGROUNDS', (0, 0), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
            ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#3182ce')),  # Professional header line
        ]
        
        # Add special styling for logo columns if present
        if has_logos and len(col_widths) > 2:
            table_style.extend([
                ('ALIGN', (2, 0), (2, -1), 'CENTER'),
                ('VALIGN', (2, 0), (2, -1), 'MIDDLE'),
                ('FONTSIZE', (2, 0), (2, -1), 8),
            ])
        
        # Apply background color to header if specified
        if bg_color:
            table_style.append(('BACKGROUND', (0, 0), (-1, 0), bg_color))
        
        table.setStyle(TableStyle(table_style))
        return table

    async def export_to_pdf(self, data: Dict[str, Any]) -> bytes:
        """Export VFX specification to professional styled PDF with enhanced visual elements"""
        try:
            logger.info("Generating enhanced professional styled PDF export")
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(
                buffer, 
                pagesize=A4, 
                topMargin=0.75*inch, 
                bottomMargin=0.75*inch,
                leftMargin=0.75*inch,
                rightMargin=0.75*inch
            )
            story = []
            
            # ENHANCED HEADER SECTION WITH PROFESSIONAL STYLING
            letterhead_info = data.get('letterheadInfo', {})
            
            # Professional header with logo and company info
            header_elements = []
            
            # Main logo placement (top center or left)
            main_logo = self._get_logo_from_data(data, 'letterheadInfo.logo')
            if main_logo:
                logo_img = self._get_logo_image(main_logo, height=1.2*inch, width=2.4*inch)
                if logo_img:
                    header_elements.append(logo_img)
                    header_elements.append(Spacer(1, 15))
            
            # Company information with enhanced styling
            if letterhead_info.get('userCompanyName'):
                company_name = Paragraph(letterhead_info['userCompanyName'], self.custom_styles['company'])
                header_elements.append(company_name)
                
                # Contact information
                contact_info = []
                if letterhead_info.get('email'):
                    contact_info.append(letterhead_info['email'])
                if letterhead_info.get('website'):
                    contact_info.append(letterhead_info['website'])
                if letterhead_info.get('address'):
                    contact_info.append(letterhead_info['address'])
                
                for info in contact_info:
                    header_elements.append(Paragraph(info, self.custom_styles['contact']))
                
                header_elements.append(Spacer(1, 20))
            
            # Add all header elements
            for element in header_elements:
                story.append(element)
            
            # ENHANCED TITLE SECTION WITH BORDERS
            title = Paragraph("IMAGE FORMAT EXCHANGE SPECS", self.custom_styles['title'])
            story.append(title)
            
            # Professional subtitle with enhanced styling
            subtitle = Paragraph("Technical Consistency Across Processes", self.custom_styles['subtitle'])
            story.append(subtitle)
            story.append(Spacer(1, 15))
            
            # Enhanced date with professional styling
            date_text = f"Document Generated: {datetime.now().strftime('%B %d, %Y at %H:%M UTC')}"
            story.append(Paragraph(date_text, self.custom_styles['date']))
            
            # Professional divider with enhanced styling
            story.append(Spacer(1, 25))
            story.append(self._create_enhanced_divider_line(colors.HexColor('#3182ce'), 3))
            story.append(Spacer(1, 30))
            
            # PROJECT INFORMATION SECTION WITH LOGO INTEGRATION
            project_info = data.get('projectInfo', {})
            if any(self._should_include_field(project_info.get(field)) for field in project_info):
                section_header = self._create_section_header("PROJECT INFORMATION", colors.HexColor('#2b6cb0'))
                story.append(KeepTogether([section_header]))
                story.append(Spacer(1, 15))
                
                # Prepare project data with enhanced organization
                project_data = []
                
                # Basic project information
                basic_fields = {
                    'documentVersion': 'Document Version:',
                    'projectDate': 'Project Date:',
                    'projectTitle': 'Project Title:',
                    'projectCodeName': 'Project Code Name:',
                    'projectFormat': 'Project Format:',
                    'projectFrameRate': 'Project Frame Rate:',
                    'colorScience': 'Color Science:',
                    'customColorScience': 'Custom Color Science:'
                }
                
                for field, label in basic_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        project_data.append([label, str(project_info[field])])
                
                if project_data:
                    table = self._create_styled_table(project_data, [2.5*inch, 3.5*inch])
                    story.append(table)
                    story.append(Spacer(1, 20))
                
                # CLIENT SECTION WITH LOGO
                client_data = []
                if self._should_include_field(project_info.get('client')):
                    client_data.append(['Client:', str(project_info['client'])])
                
                # Add client logo if available
                client_logo = self._get_logo_from_data(data, 'projectInfo.clientLogo')
                if client_logo and client_data:
                    logo_img = self._get_logo_image(client_logo, height=0.8*inch, width=1.2*inch)
                    if logo_img:
                        client_data.append(['Client Logo:', '', logo_img])
                        table = self._create_styled_table(client_data, [2*inch, 2.5*inch, 1.5*inch], has_logos=True)
                    else:
                        table = self._create_styled_table(client_data, [2.5*inch, 3.5*inch])
                elif client_data:
                    table = self._create_styled_table(client_data, [2.5*inch, 3.5*inch])
                else:
                    table = None
                
                if table:
                    story.append(table)
                    story.append(Spacer(1, 20))
                
                # PRODUCTION TEAM SECTION WITH LOGOS
                production_data = []
                production_fields = {
                    'director': 'Director:',
                    'dop': 'Director of Photography:',
                    'productionCompany': 'Production Company:'
                }
                
                for field, label in production_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        production_data.append([label, str(project_info[field])])
                
                # Add production company logo if available
                prod_logo = self._get_logo_from_data(data, 'projectInfo.productionCompanyLogo')
                if prod_logo and production_data:
                    logo_img = self._get_logo_image(prod_logo, height=0.8*inch, width=1.2*inch)
                    if logo_img:
                        production_data.append(['Production Company Logo:', '', logo_img])
                        table = self._create_styled_table(production_data, [2*inch, 2.5*inch, 1.5*inch], has_logos=True)
                    else:
                        table = self._create_styled_table(production_data, [2.5*inch, 3.5*inch])
                elif production_data:
                    table = self._create_styled_table(production_data, [2.5*inch, 3.5*inch])
                else:
                    table = None
                
                if table:
                    story.append(table)
                    story.append(Spacer(1, 20))
                
                # POST-PRODUCTION SECTION WITH LOGOS
                post_data = []
                post_fields = {
                    'postProductionSupervisor': 'Post-Production Supervisor:',
                    'lab': 'Lab:',
                    'colorist': 'Colorist:'
                }
                
                for field, label in post_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        post_data.append([label, str(project_info[field])])
                
                # Add lab logo if available
                lab_logo = self._get_logo_from_data(data, 'projectInfo.labLogo')
                if lab_logo and post_data:
                    logo_img = self._get_logo_image(lab_logo, height=0.8*inch, width=1.2*inch)
                    if logo_img:
                        post_data.append(['Lab Logo:', '', logo_img])
                        table = self._create_styled_table(post_data, [2*inch, 2.5*inch, 1.5*inch], has_logos=True)
                    else:
                        table = self._create_styled_table(post_data, [2.5*inch, 3.5*inch])
                elif post_data:
                    table = self._create_styled_table(post_data, [2.5*inch, 3.5*inch])
                else:
                    table = None
                
                if table:
                    story.append(table)
                    story.append(Spacer(1, 20))
                
                # VFX SECTION WITH VENDOR LOGO
                vfx_data = []
                vfx_fields = {
                    'vfxSupervisor': 'VFX Supervisor:',
                    'vfxOnSetSupervisor': 'VFX On-Set Supervisor:',
                    'vfxVendor': 'VFX Vendor:',
                    'vendorCodeName': 'Vendor Code Name:'
                }
                
                for field, label in vfx_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        vfx_data.append([label, str(project_info[field])])
                
                # Add VFX vendor logo if available
                vfx_logo = self._get_logo_from_data(data, 'projectInfo.vfxVendorLogo')
                if vfx_logo and vfx_data:
                    logo_img = self._get_logo_image(vfx_logo, height=0.8*inch, width=1.2*inch)
                    if logo_img:
                        vfx_data.append(['VFX Vendor Logo:', '', logo_img])
                        table = self._create_styled_table(vfx_data, [2*inch, 2.5*inch, 1.5*inch], has_logos=True)
                    else:
                        table = self._create_styled_table(vfx_data, [2.5*inch, 3.5*inch])
                elif vfx_data:
                    table = self._create_styled_table(vfx_data, [2.5*inch, 3.5*inch])
                else:
                    table = None
                
                if table:
                    story.append(table)
                    story.append(Spacer(1, 20))
                
                # Links section
                if self._should_include_field(project_info.get('vfxDocumentsLink')):
                    links_data = [['VFX Documents Link:', str(project_info['vfxDocumentsLink'])]]
                    table = self._create_styled_table(links_data, [2.5*inch, 3.5*inch])
                    story.append(table)
                    story.append(Spacer(1, 25))
                
                # Add decorative separator
                story.append(self._create_decorative_border())
                story.append(Spacer(1, 15))
            
            # CAMERA FORMATS SECTION with enhanced styling
            camera_formats = data.get('cameraFormats', [])
            if camera_formats:
                section_header = self._create_section_header("CAMERA FORMATS", colors.HexColor('#38a169'))
                story.append(KeepTogether([section_header]))
                story.append(Spacer(1, 15))
                
                for i, camera in enumerate(camera_formats, 1):
                    if any(self._should_include_field(camera.get(field)) for field in camera):
                        # Enhanced camera subsection header
                        subsection_style = ParagraphStyle(
                            'CameraSubsection',
                            fontSize=13,
                            spaceBefore=15,
                            spaceAfter=10,
                            textColor=colors.HexColor('#1a365d'),
                            fontName='Helvetica-Bold',
                            borderWidth=1,
                            borderPadding=8,
                            borderColor=colors.HexColor('#38a169'),
                            backColor=colors.HexColor('#f0fff4')
                        )
                        story.append(Paragraph(f"Camera Configuration {i}: {camera.get('cameraId', 'Unknown')}", subsection_style))
                        
                        camera_data = []
                        camera_fields = {
                            'sourceCamera': 'Source Camera:',
                            'codec': 'Codec:',
                            'sensorMode': 'Sensor Mode:',
                            'lensSqueezeeFactor': 'Lens Squeeze Factor:',
                            'colorSpace': 'Color Space/Transfer Function:'
                        }
                        
                        for field, label in camera_fields.items():
                            if self._should_include_field(camera.get(field)):
                                camera_data.append([label, str(camera[field])])
                        
                        if camera_data:
                            table = self._create_styled_table(camera_data, [2.2*inch, 3.8*inch])
                            story.append(table)
                            story.append(Spacer(1, 15))
                
                story.append(self._create_decorative_border())
                story.append(Spacer(1, 15))
            
            # VFX PULLS SECTION with enhanced purple styling
            vfx_pulls = data.get('vfxPulls', {})
            if any(self._should_include_field(vfx_pulls.get(field)) for field in vfx_pulls):
                section_header = self._create_section_header("VFX PULLS SPECIFICATIONS", colors.HexColor('#805ad5'))
                story.append(KeepTogether([section_header]))
                story.append(Spacer(1, 15))
                
                # Group VFX pulls data for better organization
                technical_data = []
                naming_data = []
                
                technical_fields = {
                    'fileFormat': 'File Format:',
                    'compression': 'Compression:',
                    'resolution': 'Resolution:',
                    'colorSpace': 'Color Space:',
                    'bitDepth': 'Bit Depth:',
                    'frameHandles': 'Frame Handles:',
                    'framePadding': 'Frame Padding:'
                }
                
                naming_fields = {
                    'showId': 'Show ID:',
                    'episode': 'Episode:',
                    'sequence': 'Sequence:',
                    'scene': 'Scene:',
                    'shotId': 'Shot ID:',
                    'plate': 'Plate:',
                    'identifier': 'Identifier:',
                    'version': 'Version:'
                }
                
                for field, label in technical_fields.items():
                    if self._should_include_field(vfx_pulls.get(field)):
                        technical_data.append([label, str(vfx_pulls[field])])
                
                for field, label in naming_fields.items():
                    if self._should_include_field(vfx_pulls.get(field)):
                        naming_data.append([label, str(vfx_pulls[field])])
                
                if technical_data:
                    # Technical specifications subsection
                    subsection_header = Paragraph("Technical Specifications", self.custom_styles['subsection'])
                    story.append(subsection_header)
                    table = self._create_styled_table(technical_data, [2.2*inch, 3.8*inch])
                    story.append(table)
                    story.append(Spacer(1, 15))
                
                if naming_data:
                    # Naming conventions subsection
                    subsection_header = Paragraph("Naming Conventions", self.custom_styles['subsection'])
                    story.append(subsection_header)
                    table = self._create_styled_table(naming_data, [2.2*inch, 3.8*inch])
                    story.append(table)
                    story.append(Spacer(1, 15))
                
                # Add VFX LUTs link if available
                if self._should_include_field(vfx_pulls.get('vfxLutsLink')):
                    link_data = [['VFX LUTs Link:', str(vfx_pulls['vfxLutsLink'])]]
                    table = self._create_styled_table(link_data, [2.2*inch, 3.8*inch])
                    story.append(table)
                    story.append(Spacer(1, 15))
                
                story.append(self._create_decorative_border())
                story.append(Spacer(1, 15))
            
            # MEDIA REVIEW SECTION with enhanced teal styling
            media_review = data.get('mediaReview', {})
            if any(self._should_include_field(media_review.get(field)) for field in media_review):
                section_header = self._create_section_header("MEDIA REVIEW SPECIFICATIONS", colors.HexColor('#319795'))
                story.append(KeepTogether([section_header]))
                story.append(Spacer(1, 15))
                
                media_data = []
                media_fields = {
                    'container': 'Container:',
                    'videoCodec': 'Video Codec:',
                    'resolution': 'Resolution:',
                    'aspectRatio': 'Aspect Ratio:',
                    'letterboxing': 'Letterboxing:',
                    'frameRate': 'Frame Rate:',
                    'colorSpace': 'Color Space:'
                }
                
                for field, label in media_fields.items():
                    if self._should_include_field(media_review.get(field)):
                        media_data.append([label, str(media_review[field])])
                
                if media_data:
                    table = self._create_styled_table(media_data, [2.2*inch, 3.8*inch])
                    story.append(table)
                    story.append(Spacer(1, 15))
                
                # Add slate & overlays link if available
                if self._should_include_field(media_review.get('slateOverlaysLink')):
                    link_data = [['Slate & Overlays Link:', str(media_review['slateOverlaysLink'])]]
                    table = self._create_styled_table(link_data, [2.2*inch, 3.8*inch])
                    story.append(table)
                    story.append(Spacer(1, 15))
                
                story.append(self._create_decorative_border())
                story.append(Spacer(1, 15))
            
            # VFX DELIVERIES SECTION with enhanced orange styling
            vfx_deliveries = data.get('vfxDeliveries', {})
            if any(self._should_include_field(vfx_deliveries.get(field)) for field in vfx_deliveries):
                section_header = self._create_section_header("VFX DELIVERIES SPECIFICATIONS", colors.HexColor('#ed8936'))
                story.append(KeepTogether([section_header]))
                story.append(Spacer(1, 15))
                
                delivery_data = []
                delivery_fields = {
                    'showId': 'Show ID:',
                    'episode': 'Episode:',
                    'sequence': 'Sequence:',
                    'scene': 'Scene:',
                    'shotId': 'Shot ID:',
                    'task': 'Task:',
                    'vendorCodeName': 'Vendor Code Name:',
                    'version': 'Version:'
                }
                
                for field, label in delivery_fields.items():
                    if self._should_include_field(vfx_deliveries.get(field)):
                        delivery_data.append([label, str(vfx_deliveries[field])])
                
                if delivery_data:
                    table = self._create_styled_table(delivery_data, [2.2*inch, 3.8*inch])
                    story.append(table)
                    story.append(Spacer(1, 25))
                
                story.append(self._create_decorative_border())
            
            # PROFESSIONAL FOOTER
            story.append(Spacer(1, 30))
            story.append(self._create_enhanced_divider_line(colors.HexColor('#2b6cb0'), 3))
            
            # Footer with document info
            footer_text = f"This document was generated automatically on {datetime.now().strftime('%B %d, %Y at %H:%M UTC')} • VFX Specifications Exchange System"
            footer_style = ParagraphStyle(
                'FooterStyle',
                fontSize=9,
                textColor=colors.HexColor('#718096'),
                alignment=1,
                fontName='Helvetica-Oblique',
                spaceBefore=10
            )
            story.append(Paragraph(footer_text, footer_style))
            
            # Build PDF with enhanced error handling
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise

    async def export_to_docx(self, data: Dict[str, Any]) -> bytes:
        """Export VFX specification to professional DOCX with enhanced styling and logo integration"""
        try:
            logger.info("Generating enhanced professional DOCX export")
            
            doc = Document()
            
            # Enhanced document styles
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)
            
            # Add custom styles for professional appearance
            heading_style = doc.styles['Heading 1']
            heading_style.font.name = 'Calibri'
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = RGBColor(26, 54, 93)  # Dark blue
            
            # ENHANCED HEADER SECTION
            letterhead_info = data.get('letterheadInfo', {})
            
            # Company information with enhanced styling
            if letterhead_info.get('userCompanyName'):
                company_para = doc.add_paragraph()
                company_run = company_para.add_run(letterhead_info['userCompanyName'])
                company_run.font.name = 'Calibri'
                company_run.font.size = Pt(20)
                company_run.font.bold = True
                company_run.font.color.rgb = RGBColor(26, 54, 93)
                company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Contact information
                contact_info = []
                if letterhead_info.get('email'):
                    contact_info.append(letterhead_info['email'])
                if letterhead_info.get('website'):
                    contact_info.append(letterhead_info['website'])
                if letterhead_info.get('address'):
                    contact_info.append(letterhead_info['address'])
                
                for info in contact_info:
                    contact_para = doc.add_paragraph(info)
                    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    contact_run = contact_para.runs[0]
                    contact_run.font.size = Pt(11)
                    contact_run.font.color.rgb = RGBColor(74, 85, 104)
                
                doc.add_paragraph()  # Empty line
            
            # Main logo if available
            main_logo = self._get_logo_from_data(data, 'letterheadInfo.logo')
            if main_logo:
                try:
                    # Extract base64 data and convert to image
                    base64_data = main_logo.split(',')[1]
                    image_data = base64.b64decode(base64_data)
                    image_buffer = io.BytesIO(image_data)
                    
                    # Add logo to document
                    logo_para = doc.add_paragraph()
                    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(image_buffer, width=Inches(3))
                    doc.add_paragraph()  # Empty line
                except Exception as e:
                    logger.warning(f"Could not add main logo to DOCX: {str(e)}")
            
            # ENHANCED TITLE SECTION
            title = doc.add_heading('IMAGE FORMAT EXCHANGE SPECS', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.runs[0]
            title_run.font.color.rgb = RGBColor(26, 54, 93)
            
            # Professional subtitle
            subtitle = doc.add_paragraph('Technical Consistency Across Processes')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.italic = True
            subtitle_run.font.color.rgb = RGBColor(74, 85, 104)
            
            # Enhanced date
            date_para = doc.add_paragraph(f"Document Generated: {datetime.now().strftime('%B %d, %Y at %H:%M UTC')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            date_run = date_para.runs[0]
            date_run.font.size = Pt(10)
            date_run.font.color.rgb = RGBColor(113, 128, 150)
            date_run.font.italic = True
            
            doc.add_paragraph()  # Empty line
            
            # Add horizontal line separator
            separator_para = doc.add_paragraph()
            separator_run = separator_para.add_run('_' * 80)
            separator_run.font.color.rgb = RGBColor(49, 130, 206)
            separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Empty line
            
            # PROJECT INFORMATION SECTION WITH ENHANCED ORGANIZATION
            project_info = data.get('projectInfo', {})
            if any(self._should_include_field(project_info.get(field)) for field in project_info):
                # Section header with enhanced styling
                section_heading = doc.add_heading('PROJECT INFORMATION', level=1)
                section_run = section_heading.runs[0]
                section_run.font.color.rgb = RGBColor(43, 108, 176)
                
                # Basic project information table
                basic_fields = {
                    'documentVersion': 'Document Version',
                    'projectDate': 'Project Date',
                    'projectTitle': 'Project Title',
                    'projectCodeName': 'Project Code Name',
                    'projectFormat': 'Project Format',
                    'projectFrameRate': 'Project Frame Rate',
                    'colorScience': 'Color Science',
                    'customColorScience': 'Custom Color Science'
                }
                
                basic_data = []
                for field, label in basic_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        basic_data.append([label, str(project_info[field])])
                
                if basic_data:
                    self._add_enhanced_docx_table(doc, basic_data, "Basic Information")
                
                # CLIENT SECTION WITH LOGO
                client_data = []
                if self._should_include_field(project_info.get('client')):
                    client_data.append(['Client', str(project_info['client'])])
                
                if client_data:
                    self._add_enhanced_docx_table(doc, client_data, "Client Information")
                    
                    # Add client logo if available
                    client_logo = self._get_logo_from_data(data, 'projectInfo.clientLogo')
                    if client_logo:
                        self._add_logo_to_docx(doc, client_logo, "Client Logo")
                
                # PRODUCTION TEAM SECTION
                production_data = []
                production_fields = {
                    'director': 'Director',
                    'dop': 'Director of Photography',
                    'productionCompany': 'Production Company'
                }
                
                for field, label in production_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        production_data.append([label, str(project_info[field])])
                
                if production_data:
                    self._add_enhanced_docx_table(doc, production_data, "Production Team")
                    
                    # Add production company logo if available
                    prod_logo = self._get_logo_from_data(data, 'projectInfo.productionCompanyLogo')
                    if prod_logo:
                        self._add_logo_to_docx(doc, prod_logo, "Production Company Logo")
                
                # POST-PRODUCTION SECTION
                post_data = []
                post_fields = {
                    'postProductionSupervisor': 'Post-Production Supervisor',
                    'lab': 'Lab',
                    'colorist': 'Colorist'
                }
                
                for field, label in post_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        post_data.append([label, str(project_info[field])])
                
                if post_data:
                    self._add_enhanced_docx_table(doc, post_data, "Post-Production")
                    
                    # Add lab logo if available
                    lab_logo = self._get_logo_from_data(data, 'projectInfo.labLogo')
                    if lab_logo:
                        self._add_logo_to_docx(doc, lab_logo, "Lab Logo")
                
                # VFX SECTION
                vfx_data = []
                vfx_fields = {
                    'vfxSupervisor': 'VFX Supervisor',
                    'vfxOnSetSupervisor': 'VFX On-Set Supervisor',
                    'vfxVendor': 'VFX Vendor',
                    'vendorCodeName': 'Vendor Code Name'
                }
                
                for field, label in vfx_fields.items():
                    if self._should_include_field(project_info.get(field)):
                        vfx_data.append([label, str(project_info[field])])
                
                if vfx_data:
                    self._add_enhanced_docx_table(doc, vfx_data, "VFX Team")
                    
                    # Add VFX vendor logo if available
                    vfx_logo = self._get_logo_from_data(data, 'projectInfo.vfxVendorLogo')
                    if vfx_logo:
                        self._add_logo_to_docx(doc, vfx_logo, "VFX Vendor Logo")
                
                # Links section
                if self._should_include_field(project_info.get('vfxDocumentsLink')):
                    links_data = [['VFX Documents Link', str(project_info['vfxDocumentsLink'])]]
                    self._add_enhanced_docx_table(doc, links_data, "Reference Links")
            
            # CAMERA FORMATS SECTION
            camera_formats = data.get('cameraFormats', [])
            if camera_formats:
                section_heading = doc.add_heading('Camera Formats', level=1)
                section_run = section_heading.runs[0]
                section_run.font.color.rgb = RGBColor(56, 161, 105)
                
                for i, camera in enumerate(camera_formats, 1):
                    if any(self._should_include_field(camera.get(field)) for field in camera):
                        camera_heading = doc.add_heading(f"Camera Configuration {i}: {camera.get('cameraId', 'Unknown')}", level=2)
                        camera_run = camera_heading.runs[0]
                        camera_run.font.color.rgb = RGBColor(26, 54, 93)
                        
                        camera_data = []
                        camera_fields = {
                            'sourceCamera': 'Source Camera',
                            'codec': 'Codec',
                            'sensorMode': 'Sensor Mode',
                            'lensSqueezeeFactor': 'Lens Squeeze Factor',
                            'colorSpace': 'Color Space/Transfer Function'
                        }
                        
                        for field, label in camera_fields.items():
                            if self._should_include_field(camera.get(field)):
                                camera_data.append([label, str(camera[field])])
                        
                        if camera_data:
                            self._add_enhanced_docx_table(doc, camera_data)
            
            # VFX PULLS SECTION
            vfx_pulls = data.get('vfxPulls', {})
            if any(self._should_include_field(vfx_pulls.get(field)) for field in vfx_pulls):
                section_heading = doc.add_heading('VFX Pulls Specifications', level=1)
                section_run = section_heading.runs[0]
                section_run.font.color.rgb = RGBColor(128, 90, 213)
                
                # Technical specifications
                technical_data = []
                technical_fields = {
                    'fileFormat': 'File Format',
                    'compression': 'Compression',
                    'resolution': 'Resolution',
                    'colorSpace': 'Color Space',
                    'bitDepth': 'Bit Depth',
                    'frameHandles': 'Frame Handles',
                    'framePadding': 'Frame Padding'
                }
                
                for field, label in technical_fields.items():
                    if self._should_include_field(vfx_pulls.get(field)):
                        technical_data.append([label, str(vfx_pulls[field])])
                
                if technical_data:
                    self._add_enhanced_docx_table(doc, technical_data, "Technical Specifications")
                
                # Naming conventions
                naming_data = []
                naming_fields = {
                    'showId': 'Show ID',
                    'episode': 'Episode',
                    'sequence': 'Sequence',
                    'scene': 'Scene',
                    'shotId': 'Shot ID',
                    'plate': 'Plate',
                    'identifier': 'Identifier',
                    'version': 'Version'
                }
                
                for field, label in naming_fields.items():
                    if self._should_include_field(vfx_pulls.get(field)):
                        naming_data.append([label, str(vfx_pulls[field])])
                
                if naming_data:
                    self._add_enhanced_docx_table(doc, naming_data, "Naming Conventions")
                
                # VFX LUTs link
                if self._should_include_field(vfx_pulls.get('vfxLutsLink')):
                    link_data = [['VFX LUTs Link', str(vfx_pulls['vfxLutsLink'])]]
                    self._add_enhanced_docx_table(doc, link_data, "Reference Links")
            
            # MEDIA REVIEW SECTION
            media_review = data.get('mediaReview', {})
            if any(self._should_include_field(media_review.get(field)) for field in media_review):
                section_heading = doc.add_heading('Media Review Specifications', level=1)
                section_run = section_heading.runs[0]
                section_run.font.color.rgb = RGBColor(49, 151, 149)
                
                media_data = []
                media_fields = {
                    'container': 'Container',
                    'videoCodec': 'Video Codec',
                    'resolution': 'Resolution',
                    'aspectRatio': 'Aspect Ratio',
                    'letterboxing': 'Letterboxing',
                    'frameRate': 'Frame Rate',
                    'colorSpace': 'Color Space'
                }
                
                for field, label in media_fields.items():
                    if self._should_include_field(media_review.get(field)):
                        media_data.append([label, str(media_review[field])])
                
                if media_data:
                    self._add_enhanced_docx_table(doc, media_data)
                
                # Slate & overlays link
                if self._should_include_field(media_review.get('slateOverlaysLink')):
                    link_data = [['Slate & Overlays Link', str(media_review['slateOverlaysLink'])]]
                    self._add_enhanced_docx_table(doc, link_data, "Reference Links")
            
            # VFX DELIVERIES SECTION
            vfx_deliveries = data.get('vfxDeliveries', {})
            if any(self._should_include_field(vfx_deliveries.get(field)) for field in vfx_deliveries):
                section_heading = doc.add_heading('VFX Deliveries Specifications', level=1)
                section_run = section_heading.runs[0]
                section_run.font.color.rgb = RGBColor(237, 137, 54)
                
                delivery_data = []
                delivery_fields = {
                    'showId': 'Show ID',
                    'episode': 'Episode',
                    'sequence': 'Sequence',
                    'scene': 'Scene',
                    'shotId': 'Shot ID',
                    'task': 'Task',
                    'vendorCodeName': 'Vendor Code Name',
                    'version': 'Version'
                }
                
                for field, label in delivery_fields.items():
                    if self._should_include_field(vfx_deliveries.get(field)):
                        delivery_data.append([label, str(vfx_deliveries[field])])
                
                if delivery_data:
                    self._add_enhanced_docx_table(doc, delivery_data)
            
            # PROFESSIONAL FOOTER
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run('_' * 80)
            footer_run.font.color.rgb = RGBColor(43, 108, 176)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            footer_text = doc.add_paragraph(f"This document was generated automatically on {datetime.now().strftime('%B %d, %Y at %H:%M UTC')} • VFX Specifications Exchange System")
            footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_text_run = footer_text.runs[0]
            footer_text_run.font.size = Pt(9)
            footer_text_run.font.color.rgb = RGBColor(113, 128, 150)
            footer_text_run.font.italic = True
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise

    def _add_enhanced_docx_table(self, doc, data, subtitle=None):
        """Add an enhanced table to DOCX with professional styling"""
        if not data:
            return
        
        if subtitle:
            subtitle_para = doc.add_paragraph(subtitle)
            subtitle_run = subtitle_para.runs[0]
            subtitle_run.font.bold = True
            subtitle_run.font.size = Pt(12)
            subtitle_run.font.color.rgb = RGBColor(45, 55, 72)
        
        # Create table
        table = doc.add_table(rows=len(data), cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Style the table
        for i, (label, value) in enumerate(data):
            row = table.rows[i]
            label_cell = row.cells[0]
            value_cell = row.cells[1]
            
            # Label cell styling
            label_para = label_cell.paragraphs[0]
            label_run = label_para.add_run(label)
            label_run.font.bold = True
            label_run.font.size = Pt(10)
            label_run.font.color.rgb = RGBColor(45, 55, 72)
            
            # Value cell styling
            value_para = value_cell.paragraphs[0]
            value_run = value_para.add_run(str(value))
            value_run.font.size = Pt(10)
            value_run.font.color.rgb = RGBColor(45, 55, 72)
            
            # Alternating row colors
            if i % 2 == 0:
                # Light background for even rows
                shading_elm_1 = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(qn('w:shd')))
                label_cell._tc.get_or_add_tcPr().append(shading_elm_1)
                shading_elm_2 = parse_xml(r'<w:shd {} w:fill="F8F9FA"/>'.format(qn('w:shd')))
                value_cell._tc.get_or_add_tcPr().append(shading_elm_2)
        
        doc.add_paragraph()  # Empty line after table

    def _add_logo_to_docx(self, doc, logo_data, caption):
        """Add a logo to DOCX document with caption"""
        try:
            # Extract base64 data and convert to image
            base64_data = logo_data.split(',')[1]
            image_data = base64.b64decode(base64_data)
            image_buffer = io.BytesIO(image_data)
            
            # Add logo to document
            logo_para = doc.add_paragraph()
            logo_run = logo_para.add_run()
            logo_run.add_picture(image_buffer, width=Inches(2))
            
            # Add caption
            caption_para = doc.add_paragraph(caption)
            caption_run = caption_para.runs[0]
            caption_run.font.size = Pt(9)
            caption_run.font.italic = True
            caption_run.font.color.rgb = RGBColor(113, 128, 150)
            
            doc.add_paragraph()  # Empty line
        except Exception as e:
            logger.warning(f"Could not add logo to DOCX: {str(e)}")
            # Add text indicating logo was not added
            error_para = doc.add_paragraph(f"[{caption} - Could not display image]")
            error_run = error_para.runs[0]
            error_run.font.size = Pt(9)
            error_run.font.color.rgb = RGBColor(185, 28, 28)

    def generate_filename(self, export_type: str, data: Dict[str, Any]) -> str:
        """Generate filename based on export type and data"""
        try:
            if export_type == 'vfxPulls':
                vfx_pulls = data.get('vfxPulls', {})
                return f"{vfx_pulls.get('showId', 'AAA')}_{vfx_pulls.get('episode', '101')}_{vfx_pulls.get('sequence', '001')}_{vfx_pulls.get('scene', '001')}_{vfx_pulls.get('shotId', '0010')}_{vfx_pulls.get('plate', 'PL01')}_{vfx_pulls.get('version', 'v001')}.{vfx_pulls.get('framePadding', '####')}.exr"
            elif export_type == 'vfxDeliveries':
                vfx_deliveries = data.get('vfxDeliveries', {})
                return f"{vfx_deliveries.get('showId', 'AAA')}_{vfx_deliveries.get('episode', '101')}_{vfx_deliveries.get('sequence', '001')}_{vfx_deliveries.get('scene', '001')}_{vfx_deliveries.get('shotId', '0010')}_{vfx_deliveries.get('task', 'comp')}_{vfx_deliveries.get('vendorCodeName', 'VEND')}_{vfx_deliveries.get('version', 'v001')}.{vfx_deliveries.get('framePadding', '####')}.exr"
            else:
                project_title = data.get('projectInfo', {}).get('projectTitle', 'VFX_Specification')
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                return f"{project_title.replace(' ', '_')}_{timestamp}"
        except Exception as e:
            logger.error(f"Error generating filename: {str(e)}")
            return f"VFX_Specification_{datetime.now().strftime('%Y%m%d_%H%M%S')}"